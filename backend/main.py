from fastapi import FastAPI, UploadFile, File, HTTPException, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse
from pydantic import BaseModel, Field
from typing import Optional, Literal, List
import os
import io
import json
from pathlib import Path
from dotenv import load_dotenv

load_dotenv()


def extract_file_content(file: UploadFile) -> str:
    """从上传的文件中提取文本内容，支持 PDF / DOCX / TXT / MD"""
    try:
        filename = file.filename or ""
        suffix = Path(filename).suffix.lower()
        content_bytes = file.file.read()
        if not content_bytes:
            return ""

        if suffix == ".pdf":
            try:
                import pdfplumber
                text_parts = []
                with pdfplumber.open(io.BytesIO(content_bytes)) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text() or ""
                        text_parts.append(page_text)
                return "\n".join(text_parts).strip()
            except Exception as e:
                # 回退到 PyPDF2
                try:
                    from PyPDF2 import PdfReader
                    reader = PdfReader(io.BytesIO(content_bytes))
                    text_parts = []
                    for page in reader.pages:
                        text_parts.append(page.extract_text() or "")
                    return "\n".join(text_parts).strip()
                except Exception:
                    return ""

        elif suffix == ".docx":
            try:
                from docx import Document
                doc = Document(io.BytesIO(content_bytes))
                paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
                # 也提取表格内容
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                paragraphs.append(cell.text.strip())
                return "\n".join(paragraphs).strip()
            except Exception:
                return ""

        elif suffix in (".doc",):
            # .doc 旧格式无法直接解析，提示用户转换
            return "[注意：检测到.doc旧格式，无法自动解析，建议转换为.docx或.pdf后重新上传，或直接将简历内容粘贴到文本框]"

        elif suffix in (".txt", ".md", ".markdown"):
            try:
                return content_bytes.decode("utf-8", errors="ignore").strip()
            except Exception:
                return content_bytes.decode("gbk", errors="ignore").strip()

        else:
            # 尝试作为文本读取
            try:
                return content_bytes.decode("utf-8", errors="ignore").strip()
            except Exception:
                return ""
    except Exception as e:
        print(f"文件解析失败 {file.filename}: {str(e)}")
        return ""

app = FastAPI(title="Agent Interviewer API", version="1.0.0")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# 挂载前端静态文件
BASE_DIR = Path(__file__).resolve().parent.parent
DIST_DIR = BASE_DIR / "dist"
if DIST_DIR.exists():
    app.mount("/assets", StaticFiles(directory=str(DIST_DIR / "assets")), name="assets")

    @app.get("/")
    async def serve_index():
        return FileResponse(str(DIST_DIR / "index.html"))

# 支持的API提供商配置
PROVIDERS = {
    "openai": {
        "base_url": "https://api.openai.com/v1",
        "models": ["gpt-4o", "gpt-4", "gpt-3.5-turbo"]
    },
    "agnes": {
        "base_url": "https://apihub.agnes-ai.com/v1",
        "models": ["agnes-2.0-flash"]
    },
    "deepseek": {
        "base_url": "https://api.deepseek.com",
        "models": ["deepseek-chat", "deepseek-coder"]
    },
    "qwen": {
        "base_url": "https://dashscope.aliyuncs.com/compatible-mode/v1",
        "models": ["qwen-turbo", "qwen-plus", "qwen-max"]
    },
    "kimi": {
        "base_url": "https://api.moonshot.cn/v1",
        "models": ["moonshot-v1-8k", "moonshot-v1-32k", "moonshot-v1-128k"]
    },
    "groq": {
        "base_url": "https://api.groq.com/openai/v1",
        "models": ["llama-3.1-70b-versatile", "mixtral-8x7b-32768", "gemma2-9b-it"]
    },
    "together": {
        "base_url": "https://api.together.xyz/v1",
        "models": ["meta-llama/Llama-3.1-70B-Instruct-Turbo", "deepseek-ai/DeepSeek-V3"]
    },
    "openrouter": {
        "base_url": "https://openrouter.ai/api/v1",
        "models": ["openai/gpt-4o", "anthropic/claude-3.5-sonnet"]
    },
    "custom": {
        "base_url": None,
        "models": ["custom-model"]
    }
}

class AnalysisRequest(BaseModel):
    jd: str
    resume: Optional[str] = None
    supplementary: Optional[str] = None
    career_plan: Optional[str] = None
    api_key: Optional[str] = None
    model: Optional[str] = "gpt-4o"
    provider: Optional[str] = "openai"
    custom_base_url: Optional[str] = None

class CompanyAnalysis(BaseModel):
    company_name: str
    industry: str
    description: str
    culture: list[str]
    values: str
    mission: str = "未找到公开信息"
    vision: str = "未找到公开信息"
    website: str = "未找到公开信息"
    business_overview: str = "未找到公开信息"
    position_in_company: str = "未找到公开信息"
    data_source: str = "基于JD内容分析"

class JobAnalysis(BaseModel):
    position: str
    department: str
    location: str
    salary: str
    job_description: str
    requirements: list[str]
    match_score: int
    match_analysis: str
    core_responsibilities: list[str] = []
    business_objectives: list[str] = []
    key_competencies: list[str] = []
    hidden_requirements: list[str] = []
    high_frequency_points: list[str] = []
    interviewer_focus: list[str] = []
    candidate_strengths: list[str] = []
    potential_risks: list[str] = []
    key_experiences: list[str] = []

class InterviewQuestion(BaseModel):
    id: int
    question: str
    category: str
    difficulty: str
    frequency: str = "中频"
    focus_point: str = ""
    suggestion: str = ""
    highlights: str = ""
    follow_up_questions: list[str] = []
    answer_suggestions: str = ""
    question_source: str = ""
    source_references: list[str] = []
    answer_references: list[str] = []
    order: int = 1

class AnalysisResponse(BaseModel):
    company_analysis: CompanyAnalysis
    job_analysis: JobAnalysis
    interview_questions: list[InterviewQuestion]
    source: str = "mock"

class ProviderInfo(BaseModel):
    name: str
    display_name: str
    base_url: Optional[str]
    models: list[str]

def get_base_url(provider: str, custom_base_url: Optional[str] = None) -> Optional[str]:
    if provider == "custom" and custom_base_url:
        return custom_base_url
    return PROVIDERS.get(provider, {}).get("base_url")

def analyze_with_llm(jd: str, resume_text: str, supplementary_text: str, career_plan: str,
                     api_key: str, model: str, provider: str, custom_base_url: str):
    try:
        from openai import OpenAI
        import json

        if not api_key:
            return None

        base_url = get_base_url(provider, custom_base_url)
        if not base_url:
            print("无效的API提供商或未提供自定义Base URL")
            return None

        client = OpenAI(api_key=api_key, base_url=base_url)

        # 截断过长的简历内容，避免超出token限制
        def _truncate(text: str, limit: int = 6000) -> str:
            if not text:
                return "未提供"
            text = text.strip()
            if len(text) > limit:
                return text[:limit] + "\n...(内容已截断)"
            return text

        resume_content = _truncate(resume_text, 6000)
        supplementary_content = _truncate(supplementary_text, 4000)
        career_plan_content = _truncate(career_plan, 2000)

        prompt = f"""
你是一位资深面试官、业务负责人及求职教练。你的任务不是生成模板化的面试答案，而是基于岗位JD、候选人简历（已解析的完整文本内容）及其他补充材料，为候选人生成最全面、最有竞争力的面试准备内容。

【输入材料】
=== 岗位JD ===
{jd}

=== 候选人简历（已从上传文件中解析的完整文本，请务必深度结合此内容进行分析） ===
{resume_content}

=== 项目作品 / GitHub / 个人网站 / 作品集 / 补充材料 ===
{supplementary_content}

=== 职业规划 ===
{career_plan_content}

⚠️ 重要：你必须深度结合上述【候选人简历】和【补充材料】的完整内容进行分析，而不是只看JD。简历深挖、项目深挖部分的问题必须基于简历中真实出现的经历、项目、公司、时间线来设计；岗位匹配度分析必须基于简历中的真实技能和经历。严禁只基于JD凭空生成。如果简历内容为空或未提供，请在相关部分明确说明"未提供简历，无法进行个性化深挖"。

请综合所有材料进行分析，不得虚构任何经历、项目、成果、数据或职责。

请按照以下JSON格式输出结果：
{{
  "company_analysis": {{
    "company_name": "从JD中提取或推测的公司名称",
    "industry": "公司所属行业",
    "description": "公司基本介绍（基于公开信息和JD内容，简洁概括）",
    "culture": ["企业文化关键词1", "企业文化关键词2", "企业文化关键词3"],
    "values": "公司核心价值观（必须基于官网或权威渠道的真实信息，如果找不到请直接填写'未找到公开信息'，严禁编造）",
    "mission": "公司使命（从官网或公开渠道获取，找不到请填写'未找到公开信息'，严禁编造）",
    "vision": "公司愿景（从官网或公开渠道获取，找不到请填写'未找到公开信息'，严禁编造）",
    "website": "公司官网完整URL链接（必须是基于公开信息的真实官网链接，如https://www.example.com；如果不确定请填写'未找到公开信息'，严禁编造不存在的链接）",
    "business_overview": "公司主营业务介绍（让候选人了解公司是做什么的、靠什么赚钱、在行业中的位置；基于公开信息，找不到请填写'未找到公开信息'）",
    "position_in_company": "该岗位在公司的定位与角色分析（结合JD说明此岗位在公司业务中扮演什么角色、对业务的价值、与哪些部门协作；帮助候选人理解自己未来的位置）",
    "data_source": "信息来源说明（如：官网、企查查、天眼查等；如果信息不充分请说明'基于JD内容推断，建议进一步查阅官网'）"
  }},
  "job_analysis": {{
    "position": "岗位名称",
    "department": "所属部门",
    "location": "工作地点",
    "salary": "薪资范围（根据市场行情推断）",
    "job_description": "岗位职责总结",
    "requirements": ["任职要求1", "任职要求2", "任职要求3"],
    "match_score": 0-100之间的匹配度分数,
    "match_analysis": "详细的匹配度分析（必须结合简历中的真实经历和技能进行对比，不要空泛）",
    "core_responsibilities": ["岗位核心职责1", "岗位核心职责2", "岗位核心职责3"],
    "business_objectives": ["业务目标1", "业务目标2"],
    "key_competencies": ["真正看重的能力1", "真正看重的能力2"],
    "hidden_requirements": ["隐性要求1", "隐性要求2"],
    "high_frequency_points": ["高频考察点1", "高频考察点2"],
    "interviewer_focus": ["面试官关注重点1", "面试官关注重点2"],
    "candidate_strengths": ["候选人与岗位的匹配优势1（结合简历）", "匹配优势2"],
    "potential_risks": ["可能存在的风险点1（结合简历）", "风险点2"],
    "key_experiences": ["哪些经历最值得重点展示1（来自简历）", "重点展示2"]
  }},
  "interview_questions": [
    {{
      "id": 1,
      "question": "预测的面试问题",
      "category": "面试板块（必须是以下之一：自我介绍、简历深挖、项目深挖、岗位能力、行为面试、案例分析、反问环节）",
      "difficulty": "难度（简单/中等/困难）",
      "frequency": "问题出现频率（高频/中频/低频）",
      "focus_point": "面试官考察点",
      "suggestion": "完整的参考回答（真实可信，符合面试交流习惯，逻辑清晰，重点突出，语言自然；必须用换行符\\n分段，每个逻辑层次单独成段，不要写成一大段）",
      "highlights": "回答亮点（专业能力、学习能力、主动性、Ownership等）",
      "follow_up_questions": ["面试官可能继续追问1", "追问2"],
      "answer_suggestions": "回答建议（注意事项、容易踩坑的地方、可以进一步优化的方向）",
      "question_source": "问题来源说明（'社媒经验贴' 或 'JD解析' 或 '简历解析' 或 '综合分析'；如果是社媒经验贴请注明平台如'牛客网面经'）",
      "source_references": ["问题来源链接1（面经/经验贴URL，可来自牛客/小红书/脉脉/知乎/CSDN/掘金/看准网等；若无可填空数组[]）"],
      "answer_references": ["回答参考链接1（可来自官方文档/官方博客，也可来自CSDN/掘金/知乎/博客园/GitHub等技术社区帖子；主观题或无强相关链接可填空数组[]）"],
      "order": 1
    }}
  ]
}}

【第一部分：公司分析要求】
- company_name / industry / description：基于JD和公开信息
- website：公司介绍自己的官方网站URL（即公司用于对外展示自身业务、文化、产品、招聘信息的官网首页，如 https://www.bytedance.com；注意：不是公司某个产品/服务的使用网站，也不是App下载页，而是公司本体介绍网站）。已知则给真实URL，找不到填'未找到公开信息'，严禁编造
- business_overview：让候选人清晰了解公司是做什么的、主营业务、盈利模式、行业地位
- position_in_company：帮助候选人理解此岗位在公司业务中的定位、扮演角色、协作关系、业务价值
- values / mission / vision：基于官网或权威渠道，找不到填'未找到公开信息'，严禁编造

【第二部分：岗位分析要求】
- 岗位核心职责：从JD中提取核心工作内容
- 业务目标：该岗位需要达成的业务成果
- 真正看重的能力：JD中隐含的核心能力要求
- 隐性要求：未明确写出但实际需要的能力
- 高频考察点：面试中大概率会问到的内容
- 面试官关注重点：面试官最想了解的方面
- 候选人与岗位的匹配优势：必须基于简历中的真实经历和技能进行对比分析
- 可能存在的风险点：简历与岗位不匹配的地方（结合简历）
- 哪些经历最值得重点展示：建议候选人重点准备的经历（必须来自简历）

【第三部分：生成面试问题要求 - 两个来源】
问题必须从两个方向综合生成：
来源一【社媒经验贴】（重要，尤其对大公司）：参考小红书、牛客网、脉脉、知乎、看准网等社媒上该公司/该岗位的面经和经验贴，预测真实出现过的面试问题。如果能找到相关面经贴的链接，必须放入 source_references 字段；小公司可能没有，可标注'JD解析'。
来源二【AI个性化解析】：基于JD + 候选人简历 + 补充材料 + 职业规划，进行个性化的问题设计，尤其是简历深挖和项目深挖部分，必须基于简历中真实出现的经历提问。

问题类型包括：
一、自我介绍：围绕岗位特点设计不同风格的问题
二、简历深挖：针对简历中每一段重要经历进行深入提问（为什么选择、职责、具体做了什么、最大挑战、如何解决、决策依据、成果衡量、优化方向、最大收获）——必须基于简历真实经历
三、项目深挖：围绕简历中的项目提出不同层级的问题（项目背景、目标、方案设计、数据分析、业务理解、产品思考、运营思考、关键决策、效果验证、经验复盘）——必须基于简历真实项目
四、岗位能力问题：结合JD生成针对性问题（专业能力、产品能力、运营能力、商业分析、行业理解、用户思维、数据分析、AI工具应用、跨部门协作、沟通能力、执行能力、学习能力、Ownership）
五、行为面试：困难经历、失败经历、压力处理、冲突处理、团队合作、主动推动、时间管理、反馈改进、学习成长、职业规划、价值观
六、案例分析：业务分析、产品优化、增长方案、运营策略、数据分析、问题拆解、优先级判断、方案设计
七、反问环节：给面试官提问的建议

【第四部分：参考答案要求】
真实性：所有答案必须严格基于候选人简历和提供的信息，不得虚构经历、数据、成果、项目、职责、技能。如材料不足，应明确指出。
表达要求（去AI化，重要）：
- 答案必须像优秀候选人在真实面试中的口语化表达，而不是AI生成的书面文章；
- 严禁使用"首先、其次、最后"、"综上所述"、"值得一提的是"、"在某种程度上"等AI套话句式；
- 严禁使用"赋能、抓手、闭环、对齐、颗粒度、打法、心智、链路"等空泛职场黑话；
- 多用短句、口语化表达，像在跟面试官聊天；
- 适当使用"其实"、"我觉得"、"当时我们"、"举个例子"等自然过渡词；
- 每个回答开头不要用"我的回答是"这种机械句式；
- 回答要有个人色彩和真实感，而不是模板化的万能答案。
竞争力优化：在不改变事实的前提下，帮助候选人展现专业能力、学习能力、主动性、Ownership、分析能力、业务理解、产品思维、运营思维、数据思维、沟通能力、执行能力、AI应用能力、复盘能力、成长速度、行业理解。
深度要求：不仅回答"做了什么"，还应体现为什么这么做、如何思考、如何决策、如何解决问题、最终效果、经验总结、能力沉淀。
参考回答分段（重要）：suggestion 字段的内容必须用换行符\\n进行分段，每个逻辑层次（背景/思路/做法/结果/总结）单独成段，便于阅读。不要把整段回答挤成一团。
回答参考链接（重要）：answer_references 字段不仅可以是官方文档/官方博客，也可以是各大技术社区、第三方平台的优质帖子，包括但不限于：CSDN、掘金、知乎、博客园、GitHub、Stack Overflow、牛客网、简书、微信公众号文章等。只要是和该问题强相关、能帮助候选人更好理解和准备的链接都可以放进来；主观行为题如确无强相关链接可填空数组[]。

【第五部分：输出格式要求】
每道题统一采用如下格式：
【问题】...
【考察点】...
【参考回答】...（口语化、去AI化）
【回答亮点】...
【面试官可能继续追问】...
【回答建议】...
【问题来源】...（社媒经验贴 / JD解析 / 简历解析 / 综合分析）
【来源链接】...（面经贴URL，若有）
【回答参考链接】...（客观题权威资料URL，若有）

【核心原则】
始终坚持：真实性＞包装；事实＞修饰；思考＞套路；自然表达＞模板答案；口语化＞书面化。帮助候选人更准确、更有竞争力地表达真实能力，而不是编造一份"完美答案"。

【频率分级】
请按照"高频（90%以上可能出现）→中频（50%~90%）→低频（拓展加分题）"三个层级组织所有问题，并优先保证高频问题回答质量。

【重要要求】
1. 【公司官网】website字段必须是公司介绍自己的官方网站（公司本体对外展示业务/文化/招聘的官网首页，如 https://www.bytedance.com），不是产品使用网站、不是App下载页、不是第三方介绍页。找不到填'未找到公开信息'，严禁编造
2. 【深度结合简历】简历深挖、项目深挖、匹配度分析必须基于简历中的真实经历和项目，严禁只基于JD凭空生成
3. 【面试问题顺序】必须按照一般面试的真实先后顺序排列，order字段从1开始递增
4. 【面试板块覆盖】必须包含以下板块，每个板块至少1-2个问题：自我介绍、简历深挖、项目深挖、岗位能力、行为面试、案例分析、反问环节
5. 【问题来源与经验贴链接】每道题必须标注 question_source（社媒经验贴/JD解析/简历解析/综合分析）；凡是来源为社媒经验贴的题目，只要能找到相关面经/经验贴链接（牛客/小红书/脉脉/知乎/看准网/CSDN/掘金等），必须放入 source_references 字段，方便候选人直接查看真实面经
6. 【回答风格】必须是口语化、去AI化的自然表达，严禁AI套话和职场黑话；suggestion必须用换行符\\n分段
7. 【回答参考链接】answer_references 可以是官方文档/博客，也可以是CSDN/掘金/知乎/博客园/GitHub/Stack Overflow等第三方社区优质帖子，只要是强相关即可；主观题无强相关链接可填空数组[]
8. 【回答建议】必须是完整的参考回答，可以直接使用，而不是'建议使用STAR法则'这种空泛建议
9. 【问题数量】至少生成10-15个问题，覆盖各个板块
10. 确保JSON格式正确，没有多余的Markdown标记
"""

        response = client.chat.completions.create(
            model=model,
            messages=[
                {"role": "system", "content": "你是一位资深面试官、业务负责人及求职教练。你的任务是基于岗位JD、候选人简历（完整文本）及其他补充材料，为候选人生成最全面、最有竞争力的面试准备内容。你必须深度结合简历真实内容进行分析，严禁只看JD。始终坚持：真实性＞包装，事实＞修饰，思考＞套路，自然口语化表达＞模板答案，去AI化。"},
                {"role": "user", "content": prompt}
            ],
            temperature=0.7,
            response_format={"type": "json_object"}
        )

        result = json.loads(response.choices[0].message.content)
        result["source"] = "llm"
        return result

    except Exception as e:
        print(f"LLM分析失败，使用模拟数据: {str(e)}")
        return None

@app.get("/")
async def root():
    return {"message": "Agent Interviewer API"}

@app.get("/api/health")
async def health_check():
    return {"status": "ok"}

@app.get("/api/providers")
async def get_providers():
    """获取所有支持的API提供商列表"""
    return {
        "providers": [
            {"name": "openai", "display_name": "OpenAI", "base_url": PROVIDERS["openai"]["base_url"], "models": PROVIDERS["openai"]["models"]},
            {"name": "agnes", "display_name": "Agnes AI (免费)", "base_url": PROVIDERS["agnes"]["base_url"], "models": PROVIDERS["agnes"]["models"]},
            {"name": "deepseek", "display_name": "DeepSeek", "base_url": PROVIDERS["deepseek"]["base_url"], "models": PROVIDERS["deepseek"]["models"]},
            {"name": "qwen", "display_name": "通义千问", "base_url": PROVIDERS["qwen"]["base_url"], "models": PROVIDERS["qwen"]["models"]},
            {"name": "kimi", "display_name": "Kimi", "base_url": PROVIDERS["kimi"]["base_url"], "models": PROVIDERS["kimi"]["models"]},
            {"name": "groq", "display_name": "Groq (快速)", "base_url": PROVIDERS["groq"]["base_url"], "models": PROVIDERS["groq"]["models"]},
            {"name": "together", "display_name": "Together AI", "base_url": PROVIDERS["together"]["base_url"], "models": PROVIDERS["together"]["models"]},
            {"name": "openrouter", "display_name": "OpenRouter", "base_url": PROVIDERS["openrouter"]["base_url"], "models": PROVIDERS["openrouter"]["models"]},
            {"name": "custom", "display_name": "自定义", "base_url": None, "models": ["custom-model"]}
        ]
    }

@app.post("/api/analyze", response_model=AnalysisResponse)
async def analyze(
    jd: str = Form(...),
    resume_file: Optional[UploadFile] = File(None),
    supplementary: str = Form(""),
    career_plan: str = Form(""),
    api_key: Optional[str] = Form(None),
    model: str = Form("gpt-4o"),
    provider: str = Form("openai"),
    custom_base_url: Optional[str] = Form(None),
    supplementary_files: List[UploadFile] = File([])
):
    if not jd.strip():
        raise HTTPException(status_code=400, detail="岗位JD不能为空")

    # 解析简历文件内容（真正读取文件文本，而非文件名）
    resume_text = ""
    if resume_file and resume_file.filename:
        resume_text = extract_file_content(resume_file)
        if not resume_text:
            resume_text = f"[简历文件 {resume_file.filename} 解析失败或为空，请将简历内容粘贴到补充材料文本框]"
        else:
            resume_text = f"[简历文件: {resume_file.filename}]\n{resume_text}"

    # 解析补充材料附件内容
    supplementary_parts = []
    if supplementary and supplementary.strip():
        supplementary_parts.append(supplementary.strip())
    if supplementary_files:
        for sf in supplementary_files:
            if sf and sf.filename:
                sf_content = extract_file_content(sf)
                if sf_content:
                    supplementary_parts.append(f"[附件: {sf.filename}]\n{sf_content}")
                else:
                    supplementary_parts.append(f"[附件: {sf.filename} 解析失败或为空]")
    supplementary_text = "\n\n".join(supplementary_parts) if supplementary_parts else ""

    # 调用 LLM 分析（传入真实文件内容）
    effective_api_key = api_key or os.getenv("OPENAI_API_KEY")
    if effective_api_key:
        llm_result = analyze_with_llm(
            jd=jd,
            resume_text=resume_text,
            supplementary_text=supplementary_text,
            career_plan=career_plan,
            api_key=effective_api_key,
            model=model,
            provider=provider,
            custom_base_url=custom_base_url or ""
        )
        if llm_result:
            return llm_result

    mock_response = {
        "company_analysis": {
            "company_name": "从JD分析的公司",
            "industry": "互联网/科技",
            "description": "根据岗位JD分析，这是一家科技公司，注重创新和技术发展。",
            "culture": ["创新", "团队合作", "客户至上"],
            "values": "未找到公开信息（演示模式，配置API Key后将从权威渠道获取真实信息）",
            "mission": "未找到公开信息（演示模式）",
            "vision": "未找到公开信息（演示模式）",
            "website": "未找到公开信息（演示模式，配置API Key后将尝试获取真实官网链接）",
            "business_overview": "未找到公开信息（演示模式，配置API Key后将分析公司主营业务）",
            "position_in_company": "未找到公开信息（演示模式，配置API Key后将分析此岗位在公司的定位）",
            "data_source": "基于JD内容推断，建议配置API Key获取权威渠道真实信息"
        },
        "job_analysis": {
            "position": "从JD提取的岗位名称",
            "department": "技术部",
            "location": "未知",
            "salary": "根据市场行情推断",
            "job_description": jd[:200] + "..." if len(jd) > 200 else jd,
            "requirements": ["具备相关技术能力", "良好的沟通能力", "团队协作精神"],
            "match_score": 75,
            "match_analysis": "基于您提供的信息，您的背景与该岗位有一定匹配度。建议重点准备技术面试，突出相关项目经验。",
            "core_responsibilities": ["负责产品设计与开发", "跨部门协作推动项目落地", "数据分析与优化"],
            "business_objectives": ["提升产品用户体验", "推动业务增长", "优化团队效率"],
            "key_competencies": ["技术能力", "产品思维", "数据分析", "沟通协作"],
            "hidden_requirements": ["快速学习能力", "抗压能力", "自驱性"],
            "high_frequency_points": ["项目经验", "技术能力", "职业规划"],
            "interviewer_focus": ["解决问题的能力", "团队协作", "成长潜力"],
            "candidate_strengths": ["相关项目经验丰富", "技术能力匹配"],
            "potential_risks": ["行业经验不足", "管理经验欠缺"],
            "key_experiences": ["XX项目经验", "XX技术栈实践"]
        },
        "interview_questions": [
            {
                "id": 1,
                "question": "请做一个简单的自我介绍",
                "category": "自我介绍",
                "difficulty": "简单",
                "frequency": "高频",
                "focus_point": "考察候选人的表达能力、核心优势提炼能力，以及与岗位的匹配度",
                "suggestion": "面试官您好，我叫XXX，毕业于XX大学XX专业。过去X年我一直在XX领域工作，主要负责XX方向。在上一段经历中，我主导了XX项目，通过XX方法解决了XX问题，最终实现了XX成果。我关注到贵公司正在招聘XX岗位，这与我的职业发展方向高度契合，希望能有机会加入团队，发挥我的专业能力。",
                "highlights": "通过具体项目成果展示专业能力，明确表达与岗位的匹配度，展现职业发展方向的一致性",
                "follow_up_questions": ["你觉得自己最大的优势是什么？", "为什么选择我们公司？", "你最近做过最有挑战性的事情是什么？"],
                "answer_suggestions": "注意控制时间在1-2分钟，突出重点，避免流水账式介绍。结合岗位JD突出相关经验，不要背诵简历。",
                "question_source": "社媒经验贴（牛客网面经）",
                "source_references": [],
                "answer_references": [],
                "order": 1
            },
            {
                "id": 2,
                "question": "请介绍一下你对这个岗位的理解，以及你具备哪些核心能力？",
                "category": "岗位能力",
                "difficulty": "中等",
                "frequency": "高频",
                "focus_point": "考察候选人对岗位的理解深度，以及核心能力与岗位需求的匹配程度",
                "suggestion": "根据我对JD的理解，这个岗位主要负责XX工作，核心要求包括XX、XX和XX能力。我具备X年相关经验，在XX技术/工具方面有深入实践。比如在上一份工作中，我使用XX完成了XX任务，提升了XX指标。我认为我的XX能力和XX经验能够很好地胜任这个岗位的要求。",
                "highlights": "展示对岗位的深入理解，用具体案例证明能力匹配，量化成果增强说服力",
                "follow_up_questions": ["你觉得这个岗位最核心的挑战是什么？", "如果让你做这个岗位，你会从哪里入手？", "你还需要提升哪些方面？"],
                "answer_suggestions": "不要只复述JD内容，要有自己的理解和分析。结合自身经验说明匹配度，避免空谈。",
                "question_source": "JD解析",
                "source_references": [],
                "answer_references": [],
                "order": 2
            },
            {
                "id": 3,
                "question": "你在简历中提到的XX项目，能详细讲讲吗？",
                "category": "简历深挖",
                "difficulty": "中等",
                "frequency": "高频",
                "focus_point": "考察项目经历的真实性、深度，以及候选人在项目中的角色和贡献",
                "suggestion": "这个项目是我在XX公司期间主导的，背景是XX业务面临XX挑战。我的角色是XX，主要负责XX。在技术方案上，我选择了XX方案，因为XX。过程中遇到了XX困难，我通过XX方法解决。最终项目交付了XX成果，数据指标提升了XX%。这个经历让我深刻理解了XX的重要性。",
                "highlights": "清晰的项目背景和目标，明确的个人角色和贡献，量化的成果数据，体现问题解决能力和技术选型能力",
                "follow_up_questions": ["项目中遇到的最大困难是什么？", "你在项目中做了哪些关键决策？", "如果重新做这个项目，你会如何优化？"],
                "answer_suggestions": "使用STAR法则（情境-任务-行动-结果），突出个人贡献而非团队成果，准备好项目细节应对追问。",
                "question_source": "简历解析",
                "source_references": [],
                "answer_references": [],
                "order": 3
            },
            {
                "id": 4,
                "question": "描述一个你遇到的技术难题及解决过程",
                "category": "项目深挖",
                "difficulty": "困难",
                "frequency": "中频",
                "focus_point": "考察技术能力、问题解决能力、学习能力和创新思维",
                "suggestion": "在XX项目中，我们遇到了XX技术难题，具体表现为XX。我首先通过XX方式分析问题根因，发现是XX导致的。然后我调研了XX、XX等几种方案，最终选择了XX方案，因为XX。实施过程中，我编写了XX代码/设计了XX架构，并通过XX测试验证。最终问题得到解决，系统性能提升了XX%。这个经历锻炼了我XX能力。",
                "highlights": "展现系统性的问题分析能力，多方案对比的决策能力，以及持续学习和创新的能力",
                "follow_up_questions": ["你尝试过哪些失败的方案？", "这个方案的优缺点是什么？", "这个问题对后续项目有什么影响？"],
                "answer_suggestions": "不要只说成功的结果，也要分享过程中的尝试和失败，展现真实的思考过程。",
                "question_source": "简历解析",
                "source_references": [],
                "answer_references": [],
                "order": 4
            },
            {
                "id": 5,
                "question": "你未来3-5年的职业规划是什么？",
                "category": "行为面试",
                "difficulty": "中等",
                "frequency": "高频",
                "focus_point": "考察职业规划清晰度、稳定性，以及与公司发展方向的匹配度",
                "suggestion": "我的职业规划分为三个阶段。短期（1年内）：快速融入团队，掌握XX技术栈，独立承担XX模块开发。中期（2-3年）：在XX领域深入钻研，成为团队的技术骨干，能够带领小组完成XX级别的项目。长期（3-5年）：向XX方向发展（技术专家/管理岗），主导XX方向的技术决策，为公司的XX业务贡献价值。贵公司的XX方向与我的规划高度契合。",
                "highlights": "展示清晰的职业发展路径，体现学习意愿和成长潜力，表达对公司的认同感",
                "follow_up_questions": ["如果公司发展方向与你的规划不一致怎么办？", "你计划如何实现这个规划？", "你最想提升的能力是什么？"],
                "answer_suggestions": "避免空谈，要有具体的目标和行动计划。结合公司实际情况，表达愿意与公司共同成长的态度。",
                "question_source": "社媒经验贴（小红书面经）",
                "source_references": [],
                "answer_references": [],
                "order": 5
            },
            {
                "id": 6,
                "question": "讲述一次你在团队中处理冲突的经历",
                "category": "行为面试",
                "difficulty": "中等",
                "frequency": "中频",
                "focus_point": "考察团队协作能力、沟通能力、冲突解决能力",
                "suggestion": "在XX项目中，我和XX同事在技术方案上产生了分歧，他主张XX方案，我倾向XX方案。我没有直接否定他，而是约他一起梳理需求，列出两个方案的优缺点对比表。通过客观分析数据，我们发现XX方案在XX方面更优。最终他认可了我的方案，我也采纳了他XX的建议。这次经历让我学会了用数据和逻辑解决分歧，而不是情绪化对抗。",
                "highlights": "展现成熟的沟通方式，用数据和事实说话，体现团队合作精神和同理心",
                "follow_up_questions": ["你当时的情绪是怎样的？", "如果对方仍然不认可怎么办？", "从这次经历中学到了什么？"],
                "answer_suggestions": "重点展示处理冲突的方法和过程，而不是批判对方。强调团队目标优先于个人观点。",
                "question_source": "综合分析",
                "source_references": [],
                "answer_references": [],
                "order": 6
            },
            {
                "id": 7,
                "question": "如果让你负责一个新的产品功能，你会如何规划和推进？",
                "category": "案例分析",
                "difficulty": "困难",
                "frequency": "中频",
                "focus_point": "考察产品思维、项目管理能力、优先级判断能力",
                "suggestion": "首先我会明确目标：这个功能要解决什么问题，目标用户是谁，成功的标准是什么。然后进行需求分析，收集用户反馈和业务方的需求，梳理出功能清单和优先级。接下来设计方案，包括技术架构、用户流程、UI设计等。在开发过程中，我会制定里程碑计划，定期同步进度，及时解决遇到的问题。最后上线后，通过数据监控和用户反馈评估效果，进行迭代优化。",
                "highlights": "展示系统化的产品思考能力，清晰的项目管理流程，数据驱动的决策方式",
                "follow_up_questions": ["如何确定需求优先级？", "如果资源不足怎么办？", "如何衡量功能的成功？"],
                "answer_suggestions": "结构清晰，逻辑完整，体现产品思维和工程思维的结合。",
                "question_source": "JD解析",
                "source_references": [],
                "answer_references": [],
                "order": 7
            },
            {
                "id": 8,
                "question": "你有什么想问我的吗？（反问环节）",
                "category": "反问环节",
                "difficulty": "简单",
                "frequency": "高频",
                "focus_point": "考察候选人的积极性、思考深度，以及对公司和岗位的真实兴趣",
                "suggestion": "我想了解三个方面：1）这个岗位入职后前三个月的主要工作目标和挑战是什么？2）团队目前在XX方向的技术栈和架构是怎样的，未来有什么演进计划？3）公司对这个岗位的考核标准和成长路径是怎样的？这些问题能帮助我更好地了解岗位期望，也能让我提前做好准备。",
                "highlights": "问题有深度，体现对岗位的认真思考，展示积极主动的态度",
                "follow_up_questions": [],
                "answer_suggestions": "不要问薪资、福利等基础问题，这些可以在后续环节了解。问能体现你对岗位和公司真正感兴趣的问题。",
                "question_source": "社媒经验贴（脉脉面经）",
                "source_references": [],
                "answer_references": [],
                "order": 8
            },
            {
                "id": 9,
                "question": "你做过最有成就感的一件事是什么？",
                "category": "行为面试",
                "difficulty": "中等",
                "frequency": "中频",
                "focus_point": "考察自我认知、成就感来源、价值观",
                "suggestion": "最有成就感的是XX项目。当时面临XX挑战，我通过XX方法解决了问题，最终实现了XX成果。这个过程中，我不仅提升了XX能力，更重要的是学会了XX思维方式。看到自己的努力带来了实实在在的价值，让我非常有成就感。",
                "highlights": "展示个人成长和价值追求，体现自我驱动力和成就感来源",
                "follow_up_questions": ["为什么这件事让你有成就感？", "这件事对你的职业发展有什么影响？", "还有哪些事情让你有成就感？"],
                "answer_suggestions": "结合职业发展，展示积极向上的价值观。",
                "question_source": "综合分析",
                "source_references": [],
                "answer_references": [],
                "order": 9
            },
            {
                "id": 10,
                "question": "如果项目进度落后，你会怎么做？",
                "category": "行为面试",
                "difficulty": "困难",
                "frequency": "低频",
                "focus_point": "考察抗压能力、问题解决能力、项目管理能力",
                "suggestion": "首先我会冷静分析原因：是需求变更、技术难题还是资源不足？然后与团队沟通，重新评估工作量和优先级，确定可以调整的部分。如果是技术难题，我会组织团队攻关或寻求外部帮助。如果是资源问题，我会与上级沟通争取更多支持。同时，我会及时向相关方同步进度变化，管理好预期。",
                "highlights": "展现冷静分析问题的能力，积极主动解决问题的态度，良好的沟通协调能力",
                "follow_up_questions": ["有没有实际经历过这种情况？", "你会如何平衡质量和进度？", "如果无法按期交付怎么办？"],
                "answer_suggestions": "展示系统性的问题解决能力，强调沟通和协作的重要性。",
                "question_source": "综合分析",
                "source_references": [],
                "answer_references": [],
                "order": 10
            }
        ],
        "source": "mock"
    }
    
    return mock_response

@app.post("/api/upload-resume")
async def upload_resume(file: UploadFile = File(...)):
    try:
        contents = await file.read()
        return {"filename": file.filename, "size": len(contents)}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"文件上传失败: {str(e)}")

@app.post("/api/test-api")
async def test_api(request: AnalysisRequest):
    """测试API Key是否有效"""
    try:
        from openai import OpenAI
        
        api_key = request.api_key or os.getenv("OPENAI_API_KEY")
        
        if not api_key:
            return {"status": "error", "message": "请提供API Key", "detail": "API Key 不能为空"}
        
        base_url = get_base_url(request.provider, request.custom_base_url)
        if not base_url:
            return {"status": "error", "message": "无效的API提供商或未提供自定义Base URL", "detail": f"提供商: {request.provider}"}
        
        client = OpenAI(api_key=api_key, base_url=base_url)
        
        response = client.chat.completions.create(
            model=request.model,
            messages=[{"role": "user", "content": "Hello"}]
        )
        
        return {
            "status": "success", 
            "message": "API Key有效",
            "detail": f"成功连接到 {request.provider}，模型: {request.model}"
        }
        
    except Exception as e:
        error_msg = str(e)
        error_type = type(e).__name__
        return {
            "status": "error", 
            "message": f"连接失败: {error_msg}",
            "detail": f"错误类型: {error_type}\n请检查 API Key、Base URL 和模型名称是否正确"
        }

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
